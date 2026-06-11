import sys
import os

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Error: python-docx is not installed. Please run: sudo apt-get install -y python3-docx")
    sys.exit(1)

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding of a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_with_spacing(doc, text="", style=None, before=0, after=6):
    """Add a paragraph with controlled spacing."""
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    return p

def create_document():
    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    
    # Configure Normal Style
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    
    # Title Page / Document Header
    p_title = add_paragraph_with_spacing(doc, before=30, after=6)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NexusHR System Documentation")
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(26)
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Deep Navy
    
    p_subtitle = add_paragraph_with_spacing(doc, before=0, after=24)
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("Local Development Setup, Architecture and API Reference Guide")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68) # Slate Gray

    p_divider = doc.add_paragraph()
    p_divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_div = p_divider.add_run("____________________________________________________")
    run_div.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    p_divider.paragraph_format.space_after = Pt(24)

    # 1. Introduction
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Project Overview")
    r_h1.bold = True
    r_h1.font.size = Pt(18)
    r_h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p_intro = add_paragraph_with_spacing(doc)
    p_intro.add_run("NexusHR is a modern, enterprise-ready Human Resource Management System (HRMS) designed to streamline employee tracking, database operations, leave management, real-time communications, performance reviews, and salary tracking. Rebranded and upgraded from the SyncWork core engine, the application is divided into a robust backend API service and an interactive frontend dashboard user interface.")

    # 2. System Architecture
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. System Architecture & Tech Stack")
    r_h2.bold = True
    r_h2.font.size = Pt(18)
    r_h2.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    p_arch = add_paragraph_with_spacing(doc)
    p_arch.add_run("The platform is engineered using a decoupled client-server architecture model:")

    p_list_be = add_paragraph_with_spacing(doc, style='List Bullet')
    r_be_title = p_list_be.add_run("Backend Engine: ")
    r_be_title.bold = True
    p_list_be.add_run("Built on Java 17 and Spring Boot 3.3.0. It leverages Spring Security (with BCrypt password hashing), Spring Data JPA for data persistence, Tomcat as the embedded container, and Spring WebSockets (STOMP protocol) for real-time channels.")

    p_list_fe = add_paragraph_with_spacing(doc, style='List Bullet')
    r_fe_title = p_list_fe.add_run("Frontend Interface: ")
    r_fe_title.bold = True
    p_list_fe.add_run("Built with React.js using TypeScript, styled using Tailwind CSS, and optimized using Vite. Communication with the backend is established via Axios HTTP calls and WebSocket protocols.")

    p_list_db = add_paragraph_with_spacing(doc, style='List Bullet')
    r_db_title = p_list_db.add_run("Database Layer: ")
    r_db_title.bold = True
    p_list_db.add_run("PostgreSQL 14+ is used as the relational database engine. Spring Data JPA with Hibernate automatically handles schema migration, table creation, and connection pooling via HikariCP.")

    # 3. Environment & Active Services
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. Local Port Mappings and Service Status")
    r_h3.bold = True
    r_h3.font.size = Pt(18)
    r_h3.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)

    table_ports = doc.add_table(rows=4, cols=4)
    table_ports.style = 'Light Shading Accent 1'
    headers_ports = ["Component", "Default Port", "Protocol / Technology", "Local Access Link"]
    
    # Style Header Row
    for idx, name in enumerate(headers_ports):
        cell = table_ports.cell(0, idx)
        cell.text = name
        set_cell_background(cell, "1A365D")
        set_cell_margins(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    ports_data = [
        ["Database Service", "5432", "PostgreSQL Relational DB", "localhost:5432"],
        ["Backend REST API", "8080", "Spring Boot / Embedded Tomcat", "http://localhost:8080/"],
        ["Frontend Dev Server", "5173", "React.js / Vite Development Server", "http://localhost:5173/"]
    ]
    
    for row_idx, data in enumerate(ports_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_ports.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(12)

    # 4. Database Setup & Seeding
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. Database Schema & Seeded Accounts")
    r_h4.bold = True
    r_h4.font.size = Pt(18)
    r_h4.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)

    p_db_info = add_paragraph_with_spacing(doc)
    p_db_info.add_run("Upon startup, Hibernate auto-generates the database tables. If the database is empty, the server automatically executes a seed operation to create three administrative accounts.")

    table_users = doc.add_table(rows=4, cols=4)
    table_users.style = 'Light Shading Accent 1'
    headers_users = ["Full Name", "Employee Code", "Department", "Primary Email"]
    
    for idx, name in enumerate(headers_users):
        cell = table_users.cell(0, idx)
        cell.text = name
        set_cell_background(cell, "1A365D")
        set_cell_margins(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    users_data = [
        ["Sahil Sepat", "SAHIL@NEXUSHR", "ADMIN", "sahil.sepat@nexushr.com"],
        ["Rudra Tiwari", "RUDRA@NEXUSHR", "ADMIN", "rudra.tiwari@nexushr.com"],
        ["Aditi Gupta", "ADITI@NEXUSHR", "ADMIN", "aditi.gupta@nexushr.com"]
    ]
    
    for row_idx, data in enumerate(users_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_users.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")

    p_spacer2 = doc.add_paragraph()
    p_spacer2.paragraph_format.space_after = Pt(12)

    # 5. REST API Documentation
    h5 = doc.add_paragraph()
    r_h5 = h5.add_run("5. Core REST API Endpoint Reference")
    r_h5.bold = True
    r_h5.font.size = Pt(18)
    r_h5.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)

    api_data = [
        ["POST", "/api/employees", "Add new employee. Generates random password & triggers async welcome email.", '{"name": "John", "empCode": "JOHN@NEXUSHR", "email": "john@example.com", "role": "USER", "department": "IT", "ctc": 80000.0}'],
        ["GET", "/api/employees", "Retrieve full list of all employees in database.", "N/A (Query String parameter allowed)"],
        ["GET", "/api/employees/code/{empCode}", "Retrieve employee details using employee code.", "N/A"],
        ["POST", "/api/employees/change-password", "Allows a user to change their login password.", '{"empCode": "SAHIL@NEXUSHR", "oldPassword": "...", "newPassword": "..."}'],
        ["PUT", "/api/employees/{id}/ctc", "Updates the CTC amount for an employee record.", "Request Body: Double numeric value (e.g. 95000.00)"],
        ["DELETE", "/api/employees/{id}", "Delete an employee and cascade deletes associated tasks, leave records, and attendance data.", "N/A"]
    ]

    table_api = doc.add_table(rows=7, cols=4)
    table_api.style = 'Light Shading Accent 1'
    headers_api = ["Method", "Endpoint Path", "Description / Behavior", "Body / Payload Example"]
    
    for idx, name in enumerate(headers_api):
        cell = table_api.cell(0, idx)
        cell.text = name
        set_cell_background(cell, "1A365D")
        set_cell_margins(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for row_idx, data in enumerate(api_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_api.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell)
            if col_idx == 0:
                # Color code GET vs POST vs DELETE/PUT
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                if text == "GET":
                    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
                elif text == "POST":
                    run.font.color.rgb = RGBColor(0x31, 0x97, 0x95)
                else:
                    run.font.color.rgb = RGBColor(0xDD, 0x6B, 0x20)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")

    p_spacer3 = doc.add_paragraph()
    p_spacer3.paragraph_format.space_after = Pt(12)

    # 6. Setup Scripts
    h6 = doc.add_paragraph()
    r_h6 = h6.add_run("6. Automated Infrastructure Scripting")
    r_h6.bold = True
    r_h6.font.size = Pt(18)
    r_h6.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)

    p_scripts = add_paragraph_with_spacing(doc)
    p_scripts.add_run("Three main bash utility scripts have been written to ensure automated local installation:")

    p_s1 = add_paragraph_with_spacing(doc, style='List Bullet')
    p_s1.add_run("setup_postgres.sh: ").bold = True
    p_s1.add_run("Installs PostgreSQL, configures the database service daemon, changes the 'postgres' superuser password to 'postgres', and checks/creates the 'nexushr' schema.")

    p_s2 = add_paragraph_with_spacing(doc, style='List Bullet')
    p_s2.add_run("setup_env.sh: ").bold = True
    p_s2.add_run("Verifies system properties, installs OpenJDK 17 and Maven, downloads local Node.js v22 binaries, extracts them to the project root, and configures matching local paths.")

    p_s3 = add_paragraph_with_spacing(doc, style='List Bullet')
    p_s3.add_run("run_local.sh: ").bold = True
    p_s3.add_run("The launcher script. Automatically checks and binds local node paths, ensures that frontend dependencies ('node_modules') are pre-installed via NPM if missing, launches the Maven Spring Boot daemon in the background (logs in backend.log), and starts the Vite development server in the foreground.")

    doc.save("NexusHR_Documentation.docx")
    print("Document successfully created: NexusHR_Documentation.docx")

if __name__ == "__main__":
    create_document()
